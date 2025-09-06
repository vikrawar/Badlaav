from pathlib import Path
from docx import Document

class FileEditor:
    def __init__(self, filepath):
        self.filepath = Path(filepath)
        if self.filepath.suffix.lower() != ".docx":
            raise ValueError("Only .docx files are supported")
        self.doc = Document(self.filepath)

    def read(self):
        return "\n".join([p.text for p in self.doc.paragraphs])

    def write(self, text):
        for p in self.doc.paragraphs:
            p.clear()
        self.doc.add_paragraph(text)
        self.doc.save(self.filepath)

    def append(self, text):
        self.doc.add_paragraph(text)
        self.doc.save(self.filepath)
